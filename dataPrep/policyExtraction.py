"""
Extract policy texts from MS word files:
1.duplication removing
2.metadata extraction [title, year, issuer, ptext]
3.save to csv files
"""
import docx, docx2txt, re, os, pickle
import pandas as pd
from win32com import client as wc

def getBailutext(filename):
    """
    get text from word files downloaded on http://www.bailuzhiku.com/
    :param filename: file path of a .docx file
    :return: a dict of structured policy data
    """
    # doc = docx.Document(filename)
    text = docx2txt.process(filename)   # convert to pure text
    fulltext = re.split('[\n]+',text)
    re_chinese = re.compile(u'[\u4e00-\u9fa5]', re.UNICODE)  # Chinese character
    re_num = re.compile(r'20[\d]+')
    textDict, contxt = {'title': '', 'year': '', 'issuer': '', 'ptext': ''}, []
    # dictionary format，contxt- policy text data
    title = filename.split('白鹿智库—')[1].rstrip('.docx')  # get title
    textDict.update({'title': title})

    # tb = doc.tables
    # for para in doc.paragraphs:
    #     fulltext.append(para.text)

    for i, itxt in enumerate(fulltext):
        if "发文机构：" in itxt:
            tmp = itxt.split('：')
            match = re.search(re_chinese, tmp[1])
            if match:
                textDict.update({'issuer': tmp[1]})  # get issuer
            else:
                textDict.update({'issuer': ''})
        if "发文字号：" in itxt:
            tmp = itxt.split('：')
            match = re.search(re_num, tmp[1])
            if match:
                textDict['year'] = match.group()    # try to get release date
        if "发布日期：" in itxt:
            tmp = itxt.split('：')
            match = re.search(re_num, tmp[1])
            if not textDict['year'] and match:
                textDict['year'] = match.group()    # try to get release date
        if "失效日期：" in itxt:
            contxt = fulltext[i+1:]
            textDict.update({'ptext': '\n'.join(contxt)})  # get raw text
            break
    return textDict


def getFabaotext(filename):
    """
    get text from word files downloaded on https://www.pkulaw.com/
    :param filename: file path
    :return: a dict of policy text
    """
    # 获取从北大法宝网站下载的文档文本
    # doc = docx.Document(filename)
    doc = docx2txt.process(filename)
    fulltext = re.split('[\n]+',doc)
    re_chinese = re.compile(u'[\u4e00-\u9fa5]', re.UNICODE)
    re_num = re.compile(r'20[\d]+')
    re_bra = re.compile(r'\(FBM.*\)')

    textDict, contxt = {'title':'','year':'','issuer':'','ptext':''}, []
    title = re.sub(r'\(FBM.*\)', '', filename)
    title = title.split('北大法宝\\')[1].strip('.docx')
    textDict.update({'title': title})
    #
    # for para in doc.paragraphs:
    #     fulltext.append(para.text)

    for i, itxt in enumerate(fulltext):
        if "发布部门" in itxt and ":" in itxt:
            tmp = itxt.split(':')
            match = re.search(re_chinese, tmp[1])
            if match:
                textDict.update({'issuer': tmp[1]})
            else:
                textDict.update({'issuer': ''})
        if "发文字号" in itxt and ":" in itxt:
            tmp = itxt.split(':')
            match = re.search(re_num, tmp[1])
            if match:
                textDict['year'] = match.group()
        if "发布日期" in itxt and ":" in itxt:
            tmp = itxt.split(':')
            match = re.search(re_num, tmp[1])
            if not textDict['year'] and match:
                textDict['year'] = match.group()
        if "法规类别" in itxt:
            contxt.extend(fulltext[i + 1:])
            textDict.update({'ptext': contxt})
            break
        elif "效力级别" in itxt:
            contxt.extend(fulltext[i + 1:])
            textDict.update({'ptext': '\n'.join(contxt)})
            break
    return textDict

def getdocPaths(rootpath):
    """trans .doc to .docx, return all filepaths"""
    if not os.path.exists(rootpath):
        print('No file available, please check the rootpath.')
        return
    filepaths = []  # filepaths of policy texts(.docx format)
    for name in os.listdir(rootpath):    # filename list in the given file path
        fn = os.path.join(rootpath, name)
        if fn.endswith('.doc'):
            if not os.path.exists(fn+'x'):     # trans doc to docx
                docTodocx(fn)
                filepaths.append(fn + 'x')
        elif fn.endswith('docx') and not name.startswith('~$') and fn not in filepaths:
            filepaths.append(fn)
    return filepaths


def docTodocx(filename):
    """convert .doc files to .docx format"""
    new_name = filename+'x'
    w = wc.Dispatch('Word.Application')
    # 或者使用下面的方法，使用启动独立的进程：
    # w = wc.DispatchEx('Word.Application')
    doc = w.Documents.Open(filename)
    doc.SaveAs(new_name, 16)
    doc.Close()
    w.Quit()


if __name__=="__main__":
    provnames = {'广东':1,'江苏':2,'山东':3,'四川':4,'内蒙古':5,'新疆':6}
    print('Which province\'s policy texts do you want to extract:\n1-广东\n2-江苏\n3-山东\n4-四川\n5-内蒙古\n6-新疆')
    namenum = int(input('>>>>>Choose a number:'))
    fnames_bailu, fnames_fabao = [],[]
    rrootp = r'D:\2020止水禪心\0.3新能源政策挖掘\数据收集'
    for it in provnames:
        if namenum == provnames.get(it):
            rootpath_bailu = rrootp+'\\文本-'+it+'\\白鹿智库'
            rootpath_fabao = rrootp + '\\文本-' + it + '\\北大法宝'
            print('路径获取中......')
            fnames_bailu = getdocPaths(rootpath_bailu)
            fnames_fabao = getdocPaths(rootpath_fabao)
            print('路径获取完毕！')
            break
    print('开始抽取政策文本......')

    nullpath, extractedTXT = [],[]   # text list
    for fpath in fnames_bailu:    #抽取白鹿智库文本
        tmp = getBailutext(fpath)
        if tmp['year']=='' or tmp['issuer']=='' or tmp['ptext']==[]:
            print(fpath+"这个文件异常！！")
            nullpath.append(fpath)
        else:
            extractedTXT.append(tmp)
    print('白鹿智库文本抽取完毕！')
    for i, fpath in enumerate(fnames_fabao):    #抽取北大法宝文本
        tmp = getFabaotext(fpath)
        if tmp['year'] =='' or  tmp['ptext']=='' or tmp['issuer']=='':
            print(fpath+"这个文件异常！！")
            nullpath.append(fpath)
        else:
            extractedTXT.append(tmp)
    print('所有文档抽取完毕！！！共有'+str(len(extractedTXT))+'个文档。')
    df = pd.DataFrame(extractedTXT)
    df.drop_duplicates()
    print(df.shape)
    print(df)

    df.to_excel(r'D:\3policyAyc\_database\_policytxt\Raw_'+list(provnames.keys())[namenum-1]+'.xlsx')

